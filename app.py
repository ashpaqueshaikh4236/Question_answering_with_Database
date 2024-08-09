import streamlit as st
from transformers import pipeline
from data_preprocessing import process_file
from database import store_user_history, phone_number_exists, add_phone_number, get_user_history
import io
from docx import Document

# Initialize Hugging Face model
qa_pipeline = pipeline("question-answering")

def save_history_as_docx(history, phone_number):
    # Create a DOCX document
    doc = Document()
    doc.add_heading(f"Chat History for {phone_number}", level=1)
    
    # Add history content
    for entry in history:
        doc.add_paragraph(f"Question: {entry.get('Question', 'N/A')}")
        doc.add_paragraph(f"Answer: {entry.get('Answer', 'N/A')}")
        doc.add_paragraph("")  # Add a line break between entries
    
    # Save DOCX to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def save_history_as_txt(history, phone_number):
    # Create a text file content
    content = f"Chat History for {phone_number}\n\n"
    for entry in history:
        content += f"Question: {entry.get('Question', 'N/A')}\n"
        content += f"Answer: {entry.get('Answer', 'N/A')}\n\n"
    
    # Save text to buffer
    buffer = io.BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return buffer.getvalue()

def main():
    st.title("Document Querying Application")

    # Main area for file upload and chat interaction
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])
    text = ""
    
    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1]
        text = process_file(uploaded_file, file_type)
        st.text_area("Extracted Text", text, height=300)
        st.session_state.extracted_text = text  # Save extracted text to session state

        # Show options to download extracted data in DOCX or TXT format
        st.sidebar.header("Download Extracted Data")
        text_file_format = st.sidebar.selectbox("Select extracted text file format:", ["DOCX", "TXT"])

        if text_file_format == "DOCX":
            doc = Document()
            doc.add_heading(f"Extracted Text", level=1)
            doc.add_paragraph(text)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.sidebar.download_button(
                label="Download Extracted Text as DOCX",
                data=buffer.getvalue(),
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif text_file_format == "TXT":
            content = f"Extracted Text\n\n{text}"
            buffer = io.BytesIO()
            buffer.write(content.encode('utf-8'))
            buffer.seek(0)
            st.sidebar.download_button(
                label="Download Extracted Text as TXT",
                data=buffer.getvalue(),
                file_name="extracted_text.txt",
                mime="text/plain"
            )

    # Sidebar for downloading history
    st.sidebar.header("Download Your History")
    verify_phone_number = st.sidebar.text_input("Enter your registered phone number to download history:", "")

    # Option to select file format for history
    history_file_format = st.sidebar.selectbox("Select history file format:", ["DOCX", "TXT"])

    if verify_phone_number:
        if phone_number_exists(verify_phone_number):
            history = get_user_history(verify_phone_number)
            if history:
                formatted_history = [{'Question': entry.get('query', 'N/A'), 'Answer': entry.get('response', 'N/A')} for entry in history]
                
                if history_file_format == "DOCX":
                    file_data = save_history_as_docx(formatted_history, verify_phone_number)
                    st.sidebar.download_button(
                        label="Download History as DOCX",
                        data=file_data,
                        file_name=f"chat_history_{verify_phone_number}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                elif history_file_format == "TXT":
                    file_data = save_history_as_txt(formatted_history, verify_phone_number)
                    st.sidebar.download_button(
                        label="Download History as TXT",
                        data=file_data,
                        file_name=f"chat_history_{verify_phone_number}.txt",
                        mime="text/plain"
                    )
            else:
                st.sidebar.write("No history available for this phone number.")
        else:
            st.sidebar.write("The phone number entered is not found in our records.")

    # Phone number input for chat
    phone_number = st.text_input("Enter your Phone Number (10 digits only) for chat", "")
    
    phone_number_message_placeholder = st.empty()
    if phone_number:
        if not phone_number.isdigit():
            phone_number_message_placeholder.error("Phone number must contain only digits.")
        elif len(phone_number) != 10:
            phone_number_message_placeholder.error("Phone number must be exactly 10 digits.")
        else:
            phone_number_message_placeholder.empty()

    # Question input for chat
    question = st.text_input("Enter your question")

    if st.button("Ask"):
        if question and phone_number:
            if phone_number.isdigit() and len(phone_number) == 10:
                if text:
                    context = text
                    result = qa_pipeline(question=question, context=context)
                    answer = result['answer']
                    
                    if phone_number_exists(phone_number):
                        store_user_history(phone_number, question, answer)
                        st.success("Your answer is below:")
                    else:
                        add_phone_number(phone_number)
                        store_user_history(phone_number, question, answer)
                        st.success("Your account has been created successfully. Your answer is below:")
                    
                    st.markdown(f"**Question:** {question}")
                    st.markdown(f"**Answer:** {answer}")
                else:
                    st.warning("No text extracted from the file. Please upload a file first.")
            else:
                st.warning("Please enter a valid phone number.")
        else:
            st.warning("Please enter a Phone Number and a question.")

    # Display stored questions and answers
    if phone_number and phone_number_exists(phone_number):
        st.subheader("Your Question and Answer History")
        history = get_user_history(phone_number)
        if history:
            formatted_history = [{'Question': entry.get('query', 'N/A'), 'Answer': entry.get('response', 'N/A')} for entry in history]
            
            for i, entry in enumerate(formatted_history):
                st.write(f"**Question {i + 1}:** {entry['Question']}")
                st.write(f"**Answer {i + 1}:** {entry['Answer']}")
                st.write("")  # Adding a blank line for spacing

if __name__ == "__main__":
    main()
